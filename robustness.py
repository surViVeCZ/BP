import re
import os
import string
import sys
import timeit
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import numpy as np

import xml_parse
import steganography
import tests


## @brief funkce mění formátování dokumentu
#@param doc_path cesta k dokumentu
#@note funkce je volána v souboru tests.py ve funkci check_robustness()
def change_font_style(doc_path):
    doc = Document(doc_path)
    new_doc = Document()
    obj_styles = new_doc.styles
    
    obj_charstyle = obj_styles.add_style('new_font_style', WD_STYLE_TYPE.CHARACTER)
    obj_font = obj_charstyle.font
    obj_font.size = Pt(14)
    obj_font.name = 'Times New Roman'

    for paragraph in doc.paragraphs:
        par = new_doc.add_paragraph()
        for run in paragraph.runs:
            words = run.text.strip().split(" ") # tady

            #kvůli zipu se nedopíše zbytek
            for word in words:
                par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                par.add_run(word +" ", style='new_font_style')

    save_path = 'robustness' 

    file = os.path.split(doc_path)
    file_name = file[1]

    try:
        full_path = os.path.join(save_path, file_name)
    except:
        print("Non existing path")
        sys.exit()
    new_doc.save(full_path)
